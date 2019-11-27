import json
from docx import Document
from docx.shared import Inches
import re
import os
import unicodedata
import requests
    
def striphtml(data):
    p = re.compile(r'<.*?>')
    return p.sub('', data)

def month(value: str) -> str:
    if(value == "01"): return "Ianuarie"
    elif(value == "02"): return "Februarie"
    elif(value == "03"): return "Martie"
    elif(value == "04"): return "Aprilie"
    elif(value == "05"): return "Mai"
    elif(value == "06"): return "Iunie"
    elif(value == "07"): return "Iulie"
    elif(value == "08"): return "August"
    elif(value == "09"): return "Septembrie"
    elif(value == "10"): return "Octombrie"
    elif(value == "11"): return "Noiembrie"
    elif(value == "12"): return "Decembrie"

def get_month(value: str) -> str:
    value = re.findall("-([0-9]*)-", value)
    return month(value[0])

def get_year(value: str) -> str:
    value = re.findall("([0-9][0-9][0-9][0-9])-", value)
    return value[0]

def strip_accents(text):

    try:
        text = unicode(text, 'utf-8')
    except NameError:
        pass

    text = unicodedata.normalize('NFD', text)\
           .encode('ascii', 'ignore')\
           .decode("utf-8")

    return str(text)

wp_json = 'http://localhost/asmi-local-dev/wp-json/wp/v2/posts/?per_page=100&page='
for i in range(1,5):
    with requests.Session() as s:
        resp = s.get(wp_json+str(i))
    postJson = resp.json()

    utfRep = {
        "\xc4\x83": "a",
        "\xc4\x82": "A",
        "\xc3\xa2": "a",
        "\xc3\x82": "A",
        "\xc3\xae": "i",
        "\xc3\x8e": "I",
        "\xc8\x99": "s",
        "\xc8\x98": "S",
        "\xc5\x9f": "s",
        "\xc5\x9e": "S",
        "\xc8\x9b": "t",
        "\xc8\x9a": "T",
        "\xc5\xa3": "t",
        "\xc5\xa2": "T"
    }
    
    for item in postJson:
        path = "C:/WORK/Site ASMI/CONTENT/Posts/"
        post = {
            "title": strip_accents(item["title"]["rendered"]),
            "content": striphtml(str(item["content"]["rendered"].encode())),
            "month": get_month(item["date"]),
            "year": get_year(item["date"])
        }
        for item in utfRep:
            post["content"].replace(item, utfRep[item])
        
        replaceArray = ["<",">","/","\\","?","*",":","|","\""]
        for item in replaceArray:
            post["title"] = post["title"].replace(item, "")
        
        document = Document()
        document.add_heading(post["title"], 0)
        p = document.add_paragraph(post["content"])

        filename = post["title"]+".docx"
        if not os.path.exists(path+post["year"]+"/"+post["month"]):
            os.makedirs(path+post["year"]+"/"+post["month"])
        filepath = path+post["year"]+"/"+post["month"]+ "/" + filename
        document.save(filepath)

        print(post["title"]+"_"+post["month"]+"-"+post["year"]+ " ---CREATED..." + "\n")
  
